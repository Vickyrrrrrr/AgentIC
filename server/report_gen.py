"""
AgentIC Report Generator
Generates PDF and DOCX reports for each build stage and full build summary.

Uses:
  - PyMuPDF (fitz)  — already in Docker image — for PDF generation
  - python-docx     — already in Docker image — for DOCX generation
"""
import io
import json
from datetime import datetime
from typing import Any, Dict, List, Optional

# ─── DOCX ────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ─── PDF via PyMuPDF ─────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

# ─── Constants ───────────────────────────────────────────────────────
BRAND_ORANGE = (0.96, 0.47, 0.13)   # #F5781F  (AgentIC brand)
BRAND_DARK   = (0.05, 0.05, 0.05)   # near-black
PAGE_W, PAGE_H = 595, 842            # A4 in points


# ═══════════════════════════════════════════════════════════════════════
#  PDF helpers (PyMuPDF)
# ═══════════════════════════════════════════════════════════════════════

def _pdf_color(page, r, g, b):
    return (r, g, b)


def _draw_pdf_header(page, title: str, subtitle: str, design_name: str,
                     timestamp: str) -> float:
    """Draw a branded header and return the y-cursor after it."""
    # Orange banner
    page.draw_rect(fitz.Rect(0, 0, PAGE_W, 70), color=None,
                   fill=BRAND_ORANGE, overlay=True)
    # Title
    page.insert_text((36, 30), "AgentIC", fontsize=18, fontname="helv",
                     color=(1, 1, 1))
    page.insert_text((36, 52), "Autonomous Silicon Studio", fontsize=9,
                     fontname="helv", color=(1, 1, 1))
    # Right-align design name
    page.insert_text((PAGE_W - 200, 30), design_name, fontsize=11,
                     fontname="helv", color=(1, 1, 1))
    page.insert_text((PAGE_W - 200, 48), timestamp, fontsize=8,
                     fontname="helv", color=(1, 1, 1))

    # Title block
    y = 90
    page.insert_text((36, y), title, fontsize=16, fontname="helv-b",
                     color=BRAND_DARK)
    y += 22
    page.insert_text((36, y), subtitle, fontsize=10, fontname="helv",
                     color=(0.4, 0.4, 0.4))
    y += 18
    # Divider
    page.draw_line(fitz.Point(36, y), fitz.Point(PAGE_W - 36, y),
                   color=BRAND_ORANGE, width=1.2)
    return y + 14


def _pdf_section(page, y: float, heading: str) -> float:
    """Draw a section heading, return new y."""
    if y > PAGE_H - 80:  # near bottom — caller handles new page
        return y
    y += 8
    page.insert_text((36, y), heading, fontsize=11, fontname="helv-b",
                     color=BRAND_ORANGE)
    y += 14
    page.draw_line(fitz.Point(36, y), fitz.Point(PAGE_W - 36, y),
                   color=(0.85, 0.85, 0.85), width=0.5)
    return y + 6


def _pdf_body(page, y: float, text: str, font="helv", size=9.5,
              color=(0.1, 0.1, 0.1), indent=36, max_w=None) -> float:
    """Insert wrapped body text, return new y. Creates extra pages via doc if needed."""
    if not text:
        return y
    max_w = max_w or (PAGE_W - indent - 36)
    # Simple word wrap
    words = text.split()
    line, lines = [], []
    est_char = max_w / (size * 0.55)
    for w in words:
        if sum(len(x) + 1 for x in line) + len(w) > est_char:
            lines.append(" ".join(line))
            line = [w]
        else:
            line.append(w)
    if line:
        lines.append(" ".join(line))

    for ln in lines:
        if y > PAGE_H - 50:
            return y  # caller must handle overflow
        page.insert_text((indent, y), ln, fontsize=size,
                         fontname=font, color=color)
        y += size + 3
    return y + 2


# ═══════════════════════════════════════════════════════════════════════
#  Stage Report — PDF
# ═══════════════════════════════════════════════════════════════════════

def generate_stage_report_pdf(payload: dict, design_name: str) -> bytes:
    """Return PDF bytes for a single stage completion report."""
    if not _PDF_OK:
        raise RuntimeError("PyMuPDF not installed")

    stage_name   = payload.get("stage_name", "UNKNOWN")
    summary      = payload.get("summary", "No summary available.")
    next_preview = payload.get("next_stage_preview", "")
    artifacts    = payload.get("artifacts", [])
    decisions    = payload.get("decisions", [])
    warnings     = payload.get("warnings", [])
    timestamp    = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    doc  = fitz.open()
    page = doc.new_page(width=PAGE_W, height=PAGE_H)

    title    = f"Stage Report — {stage_name.replace('_', ' ').title()}"
    subtitle = f"Design: {design_name}  |  Generated: {timestamp}"

    y = _draw_pdf_header(page, title, subtitle, design_name, timestamp)

    # Summary
    y = _pdf_section(page, y, "Stage Summary")
    y = _pdf_body(page, y, summary)

    # Next stage preview
    if next_preview:
        y = _pdf_section(page, y, "Next Stage Preview")
        y = _pdf_body(page, y, next_preview)

    # Artifacts
    if artifacts:
        y = _pdf_section(page, y, "Artifacts Produced")
        for a in artifacts:
            y = _pdf_body(page, y,
                          f"• {a['name']}: {a['description']}",
                          font="helv-b", size=9)
            path = a.get("path", "")
            if path and not path.startswith("{"):
                y = _pdf_body(page, y, f"  Path: {path[:120]}",
                              color=(0.4, 0.4, 0.4), size=8.5, indent=46)

    # Decisions
    if decisions:
        y = _pdf_section(page, y, "Autonomous Decisions")
        for d in decisions:
            y = _pdf_body(page, y, f"• {d}", size=9)

    # Warnings
    if warnings:
        y = _pdf_section(page, y, "Warnings")
        for w in warnings:
            y = _pdf_body(page, y, f"⚠ {w}", size=9, color=(0.7, 0.4, 0.0))

    # Footer
    page.draw_line(fitz.Point(36, PAGE_H - 30),
                   fitz.Point(PAGE_W - 36, PAGE_H - 30),
                   color=(0.85, 0.85, 0.85), width=0.5)
    page.insert_text((36, PAGE_H - 18),
                     "Generated by AgentIC — Autonomous Silicon Studio",
                     fontsize=7.5, fontname="helv", color=(0.6, 0.6, 0.6))

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
#  Full Build Report — PDF
# ═══════════════════════════════════════════════════════════════════════

def generate_full_report_pdf(stages: Dict[str, dict], design_name: str,
                             build_status: str, events: List[dict]) -> bytes:
    """Return PDF bytes for the full build report (all stages)."""
    if not _PDF_OK:
        raise RuntimeError("PyMuPDF not installed")

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M UTC")
    doc  = fitz.open()
    page = doc.new_page(width=PAGE_W, height=PAGE_H)

    title    = f"Full Build Report — {design_name}"
    subtitle = f"Status: {build_status}  |  Generated: {timestamp}"

    y = _draw_pdf_header(page, title, subtitle, design_name, timestamp)

    def _new_page():
        nonlocal page, y
        page = doc.new_page(width=PAGE_W, height=PAGE_H)
        # Simple continuation header
        page.draw_rect(fitz.Rect(0, 0, PAGE_W, 28), color=None,
                       fill=BRAND_ORANGE, overlay=True)
        page.insert_text((36, 18), f"AgentIC — {design_name} Build Report (cont.)",
                         fontsize=8.5, fontname="helv", color=(1, 1, 1))
        y = 44

    def _ensure_space(needed=60):
        if y > PAGE_H - needed:
            _new_page()

    # Build overview
    y = _pdf_section(page, y, "Build Overview")
    total_stages = len(stages)
    y = _pdf_body(page, y, f"Design: {design_name}")
    y = _pdf_body(page, y, f"Status: {build_status}")
    y = _pdf_body(page, y, f"Total stages completed: {total_stages}")
    y = _pdf_body(page, y, f"Report generated: {timestamp}")

    # Per-stage summaries
    from server.stage_summary import STAGE_FLOW
    for stage_name in STAGE_FLOW:
        if stage_name not in stages:
            continue
        payload = stages[stage_name]
        _ensure_space(100)
        y = _pdf_section(page, y, f"Stage: {stage_name.replace('_', ' ').title()}")
        summary = payload.get("summary", "No summary.")
        y = _pdf_body(page, y, summary)

        artifacts = payload.get("artifacts", [])
        if artifacts:
            y = _pdf_body(page, y, "Artifacts:", font="helv-b", size=9)
            for a in artifacts:
                _ensure_space(30)
                y = _pdf_body(page, y, f"  • {a['name']}: {a['description']}",
                              size=8.5)

        decisions = payload.get("decisions", [])
        if decisions:
            y = _pdf_body(page, y, "Decisions:", font="helv-b", size=9)
            for d in decisions:
                _ensure_space(20)
                y = _pdf_body(page, y, f"  • {d}", size=8.5)

        warnings = payload.get("warnings", [])
        if warnings:
            for w in warnings:
                _ensure_space(20)
                y = _pdf_body(page, y, f"  ⚠ {w}", size=8.5,
                              color=(0.7, 0.4, 0.0))

    # Footer on last page
    page.draw_line(fitz.Point(36, PAGE_H - 30),
                   fitz.Point(PAGE_W - 36, PAGE_H - 30),
                   color=(0.85, 0.85, 0.85), width=0.5)
    page.insert_text((36, PAGE_H - 18),
                     "Generated by AgentIC — Autonomous Silicon Studio",
                     fontsize=7.5, fontname="helv", color=(0.6, 0.6, 0.6))

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
#  Stage Report — DOCX
# ═══════════════════════════════════════════════════════════════════════

def _docx_heading(doc, text: str, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xF5, 0x78, 0x1F)
    return p


def _docx_para(doc, text: str, bold=False, italic=False,
               color: Optional[RGBColor] = None, size_pt=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return p


def generate_stage_report_docx(payload: dict, design_name: str) -> bytes:
    """Return DOCX bytes for a single stage completion report."""
    if not _DOCX_OK:
        raise RuntimeError("python-docx not installed")

    stage_name   = payload.get("stage_name", "UNKNOWN")
    summary      = payload.get("summary", "No summary available.")
    next_preview = payload.get("next_stage_preview", "")
    artifacts    = payload.get("artifacts", [])
    decisions    = payload.get("decisions", [])
    warnings     = payload.get("warnings", [])
    timestamp    = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    doc = Document()

    # Title
    t = doc.add_heading(
        f"AgentIC Stage Report — {stage_name.replace('_', ' ').title()}", 0)
    for run in t.runs:
        run.font.color.rgb = RGBColor(0xF5, 0x78, 0x1F)

    doc.add_paragraph(f"Design: {design_name}  |  Generated: {timestamp}")
    doc.add_paragraph()

    # Summary
    _docx_heading(doc, "Stage Summary")
    doc.add_paragraph(summary)

    # Next stage
    if next_preview:
        _docx_heading(doc, "Next Stage Preview")
        doc.add_paragraph(next_preview)

    # Artifacts
    if artifacts:
        _docx_heading(doc, "Artifacts Produced")
        for a in artifacts:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{a['name']}: ")
            run.bold = True
            p.add_run(a["description"])
            path = a.get("path", "")
            if path and not path.startswith("{"):
                doc.add_paragraph(f"   Path: {path[:200]}").runs[0].font.size = Pt(8)

    # Decisions
    if decisions:
        _docx_heading(doc, "Autonomous Decisions")
        for d in decisions:
            doc.add_paragraph(d, style="List Bullet")

    # Warnings
    if warnings:
        _docx_heading(doc, "Warnings")
        for w in warnings:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"⚠ {w}")
            run.font.color.rgb = RGBColor(0xB3, 0x66, 0x00)

    # Footer
    doc.add_paragraph()
    _docx_para(doc, "Generated by AgentIC — Autonomous Silicon Studio",
               italic=True, color=RGBColor(0x80, 0x80, 0x80), size_pt=8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
#  Full Build Report — DOCX
# ═══════════════════════════════════════════════════════════════════════

def generate_full_report_docx(stages: Dict[str, dict], design_name: str,
                               build_status: str, events: List[dict]) -> bytes:
    """Return DOCX bytes for the full multi-stage build report."""
    if not _DOCX_OK:
        raise RuntimeError("python-docx not installed")

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M UTC")
    doc = Document()

    # Cover
    t = doc.add_heading(f"AgentIC Full Build Report", 0)
    for run in t.runs:
        run.font.color.rgb = RGBColor(0xF5, 0x78, 0x1F)

    doc.add_paragraph(f"Design: {design_name}")
    doc.add_paragraph(f"Status: {build_status}")
    doc.add_paragraph(f"Generated: {timestamp}")
    doc.add_paragraph()

    # Overview table
    _docx_heading(doc, "Build Overview")
    from server.stage_summary import STAGE_FLOW
    total_stages = len(stages)
    doc.add_paragraph(f"Completed stages: {total_stages} / {len(STAGE_FLOW) - 1}")
    doc.add_paragraph()

    # Per-stage sections
    for stage_name in STAGE_FLOW:
        if stage_name not in stages:
            continue
        payload = stages[stage_name]

        doc.add_page_break()
        h = doc.add_heading(f"Stage: {stage_name.replace('_', ' ').title()}", 1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xF5, 0x78, 0x1F)

        summary = payload.get("summary", "No summary.")
        doc.add_paragraph(summary)

        artifacts = payload.get("artifacts", [])
        if artifacts:
            _docx_heading(doc, "Artifacts", level=2)
            for a in artifacts:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{a['name']}: ")
                run.bold = True
                p.add_run(a["description"])
                path = a.get("path", "")
                if path and not path.startswith("{"):
                    doc.add_paragraph(f"   Path: {path[:200]}").runs[0].font.size = Pt(8)

        decisions = payload.get("decisions", [])
        if decisions:
            _docx_heading(doc, "Autonomous Decisions", level=2)
            for d in decisions:
                doc.add_paragraph(d, style="List Bullet")

        warnings = payload.get("warnings", [])
        if warnings:
            _docx_heading(doc, "Warnings", level=2)
            for w in warnings:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"⚠ {w}")
                run.font.color.rgb = RGBColor(0xB3, 0x66, 0x00)

    doc.add_paragraph()
    _docx_para(doc, "Generated by AgentIC — Autonomous Silicon Studio",
               italic=True, color=RGBColor(0x80, 0x80, 0x80), size_pt=8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
